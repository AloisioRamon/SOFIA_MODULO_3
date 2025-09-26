import streamlit as st
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
import re
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# ---------- Funções auxiliares ----------
def limpar_resposta(text: str) -> str:
    """Remove o bloco <think>...</think> do texto gerado pela IA"""
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()

def gerar_pdf(conteudo: str) -> BytesIO:
    """Gera um PDF a partir do texto"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = [Paragraph(conteudo, styles["Normal"])]
    doc.build(story)
    buffer.seek(0)
    return buffer

def gerar_docx(conteudo: str) -> BytesIO:
    """Gera um DOCX a partir do texto"""
    buffer = BytesIO()
    doc = Document()
    doc.add_paragraph(conteudo)
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extrair_texto_docx(file: BytesIO) -> str:
    """Extrai texto de um arquivo DOCX"""
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def extrair_texto_pdf(file: BytesIO) -> str:
    """Extrai texto de um arquivo PDF"""
    reader = PdfReader(file)
    texto = ""
    for page in reader.pages:
        texto += page.extract_text() + "\n"
    return texto

# ---------- Função principal ----------
def show():
    st.title("📝 Geração de Conteúdo com IA")

    # Entrada da chave da API Groq
    api_key = st.text_input("Digite sua chave da Groq", type="password")

    # Seleção do modelo
    modelo = st.selectbox(
        "Escolha o modelo",
        ["deepseek-r1-distill-llama-70b", "llama3-70b-8192"]
    )

    # Seleção do idioma
    idioma = st.selectbox(
        "Escolha o idioma do conteúdo",
        ["Português", "中文 (Chinês)", "English", "Español"]
    )

    # Upload do documento
    st.subheader("📂 Envie um documento para a IA ler")
    arquivo = st.file_uploader("Escolha um arquivo DOCX ou PDF", type=["docx", "pdf"])

    # Botão para processar documento
    if st.button("▶️ Rodar Documento"):
        if not arquivo:
            st.warning("❌ Por favor, envie um arquivo primeiro.")
        elif not api_key:
            st.error("❌ Você precisa informar a chave da Groq.")
        else:
            # Extrair texto do arquivo
            texto_documento = ""
            if arquivo.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                texto_documento = extrair_texto_docx(arquivo)
            elif arquivo.type == "application/pdf":
                texto_documento = extrair_texto_pdf(arquivo)

            if not texto_documento.strip():
                st.warning("⚠️ O arquivo está vazio ou não foi possível extrair o texto.")
            else:
                try:
                    llm = ChatGroq(model=modelo, api_key=api_key)

                    # Prompt baseado no idioma
                    if idioma == "Português":
                        instrucao = "Crie um conteúdo didático em português com base no seguinte texto: {entrada}."
                    elif idioma == "中文 (Chinês)":
                        instrucao = "请用中文撰写一个教学内容，内容如下: {entrada}。"
                    elif idioma == "English":
                        instrucao = "Create an educational content in English based on the following text: {entrada}."
                    elif idioma == "Español":
                        instrucao = "Cree un contenido educativo en español basado en el siguiente texto: {entrada}."

                    prompt = ChatPromptTemplate.from_messages([
                        ("system", "Você é um professor especialista em educação."),
                        ("user", instrucao)
                    ])

                    chain = prompt | llm | StrOutputParser()

                    with st.spinner("Gerando conteúdo do documento..."):
                        resultado = chain.invoke({"entrada": texto_documento})
                        resultado = limpar_resposta(resultado)

                    st.success("✅ Conteúdo gerado com sucesso!")
                    st.write("### Resultado")
                    st.write(resultado)

                    # Botões de download
                    pdf_buffer = gerar_pdf(resultado)
                    docx_buffer = gerar_docx(resultado)

                    st.download_button(
                        label="📄 Baixar em PDF",
                        data=pdf_buffer,
                        file_name="documento_gerado.pdf",
                        mime="application/pdf"
                    )

                    st.download_button(
                        label="📝 Baixar em Word",
                        data=docx_buffer,
                        file_name="documento_gerado.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"Erro ao gerar conteúdo: {e}")

